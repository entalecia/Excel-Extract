import pandas as pd
import docx


# f = open("file.txt", "w")
f = docx.Document()

file = pd.read_excel("./doc.xlsx")

titles = file.columns.to_list()
titles_to_ignore = ["Timestamp"]

for i, rows in file.iterrows():
    for title_name in titles:
        if title_name.title() in titles_to_ignore:
            continue
        p = f.add_paragraph()
        instance = p.add_run(f"{str(title_name)}: ")
        instance.bold = True
        p.add_run(f"{str(rows[title_name]).title()}")
    f.add_page_break()
    

f.save('file.docx')
    
    